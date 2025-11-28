#!/usr/bin/env python3
# coding: utf-8
"""
Renomeador e Conversor → PDF (com backup, progresso, relatório Excel)
- Nome final = nome_base (+ _keyword se existir)
- Se já existir, adiciona suffix _1, _2, ...
- Pede pasta destino ANTES do processamento
- Relatório Excel com colunas escolhidas (inclui palavra-chave)
"""

import os
import re
import shutil
import threading
import time
import traceback
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from fpdf import FPDF
import openpyxl
from docx import Document
from datetime import datetime

# ========== CONFIG ==========
# ajuste conforme seu sistema se necessário:
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
CAMINHO_POPPLER = r"C:\poppler-25.11.0\Library\bin"
# ===========================

# -------- utils ----------
def safe_read_text(path):
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""

def generate_unique_path(path):
    base, ext = os.path.splitext(path)
    counter = 1
    new = path
    while os.path.exists(new):
        new = f"{base}_{counter}{ext}"
        counter += 1
    return new

def ensure_dir(p):
    if not os.path.exists(p):
        os.makedirs(p, exist_ok=True)

# -------- extratores de "palavra-chave" ----------
def buscar_chave_texto(texto, palavra_chave):
    if not palavra_chave:
        return None
    padrao = rf"{re.escape(palavra_chave)}[: ]+([A-Za-z0-9\-\._]+)"
    m = re.search(padrao, texto, re.IGNORECASE)
    return m.group(1) if m else None

def extrair_chave_pdf(path, palavra_chave=None):
    try:
        pages = convert_from_path(path, poppler_path=CAMINHO_POPPLER)
        texto = ""
        for p in pages:
            texto += pytesseract.image_to_string(p, lang="por")
        return buscar_chave_texto(texto, palavra_chave)
    except Exception:
        return None

def extrair_chave_imagem(path, palavra_chave=None):
    try:
        img = Image.open(path)
        texto = pytesseract.image_to_string(img, lang="por")
        return buscar_chave_texto(texto, palavra_chave)
    except Exception:
        return None

def extrair_chave_txt(path, palavra_chave=None):
    texto = safe_read_text(path)
    return buscar_chave_texto(texto, palavra_chave)

def extrair_chave_docx(path, palavra_chave=None):
    try:
        doc = Document(path)
        texto = "\n".join([p.text for p in doc.paragraphs])
        return buscar_chave_texto(texto, palavra_chave)
    except Exception:
        return None

# -------- conversor para PDF ----------
def converter_para_pdf(caminho_arquivo, caminho_destino):
    ext = os.path.splitext(caminho_arquivo)[1].lower()
    try:
        if ext == ".pdf":
            shutil.copy2(caminho_arquivo, caminho_destino)
            return True
        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"):
            img = Image.open(caminho_arquivo).convert("RGB")
            img.save(caminho_destino)
            return True
        if ext == ".txt":
            texto = safe_read_text(caminho_arquivo)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=12)
            pdf.set_font("Arial", size=11)
            for line in texto.splitlines():
                pdf.multi_cell(0, 6, line)
            pdf.output(caminho_destino)
            return True
        if ext == ".docx":
            doc = Document(caminho_arquivo)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=12)
            pdf.set_font("Arial", size=11)
            for p in doc.paragraphs:
                for ln in (p.text or "").splitlines():
                    pdf.multi_cell(0, 6, ln)
            pdf.output(caminho_destino)
            return True
        # outros arquivos: cria PDF placeholder com nome/caminho
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, f"Arquivo original: {os.path.basename(caminho_arquivo)}")
        pdf.multi_cell(0, 6, f"Caminho original: {caminho_arquivo}")
        pdf.output(caminho_destino)
        return True
    except Exception as e:
        print(f"Erro converter {caminho_arquivo}: {e}")
        traceback.print_exc()
        return False

# ========= APP ============
class App:
    def __init__(self, root):
        self.root = root
        root.title("Renomeador Avançado → PDF")
        root.geometry("1100x700")

        # vars
        self.orig_folder = tk.StringVar()
        self.dest_folder = tk.StringVar()
        self.palavra_chave = tk.StringVar()
        self.nome_base = tk.StringVar()
        self.ext_filtro = tk.StringVar()
        self.backup_var = tk.BooleanVar(value=False)

        self.registros = []  # lista de dicts
        self.thread = None
        self.cancel_flag = False

        # UI top
        top = tk.Frame(root)
        top.pack(fill="x", padx=8, pady=6)

        tk.Label(top, text="Pasta Origem:").grid(row=0, column=0, sticky="w")
        tk.Entry(top, textvariable=self.orig_folder, width=70).grid(row=0, column=1, sticky="w")
        tk.Button(top, text="Selecionar Origem", command=self.selecionar_origem).grid(row=0, column=2, padx=6)

        tk.Label(top, text="Pasta Destino:").grid(row=1, column=0, sticky="w")
        tk.Entry(top, textvariable=self.dest_folder, width=70).grid(row=1, column=1, sticky="w")
        tk.Button(top, text="Selecionar Destino", command=self.selecionar_destino).grid(row=1, column=2, padx=6)

        tk.Label(top, text="Nome base (obrigatório):").grid(row=2, column=0, sticky="w")
        tk.Entry(top, textvariable=self.nome_base, width=30).grid(row=2, column=1, sticky="w")
        tk.Label(top, text="Palavra-chave (opcional):").grid(row=2, column=2, sticky="w")
        tk.Entry(top, textvariable=self.palavra_chave, width=30).grid(row=2, column=3, sticky="w")

        tk.Label(top, text="Filtrar extensão (opcional):").grid(row=3, column=0, sticky="w")
        tk.Entry(top, textvariable=self.ext_filtro, width=20).grid(row=3, column=1, sticky="w")
        tk.Checkbutton(top, text="Salvar BACKUP na origem", variable=self.backup_var).grid(row=3, column=2, sticky="w")
        tk.Button(top, text="Atualizar Lista", command=self.atualizar_lista).grid(row=3, column=3, sticky="w")

        # treeview
        cols = ("orig", "novo", "status")
        self.tree = ttk.Treeview(root, columns=cols, show="headings", height=22)
        self.tree.heading("orig", text="Nome Original")
        self.tree.heading("novo", text="Novo Nome")
        self.tree.heading("status", text="Status")
        self.tree.column("orig", width=420)
        self.tree.column("novo", width=420)
        self.tree.column("status", width=150)
        self.tree.pack(fill="both", expand=True, padx=8, pady=8)

        # progress + buttons
        bottom = tk.Frame(root)
        bottom.pack(fill="x", padx=8, pady=6)

        self.progress = ttk.Progressbar(bottom, orient="horizontal", mode="determinate")
        self.progress.pack(side="left", fill="x", expand=True, padx=(0,6))

        btns = tk.Frame(bottom)
        btns.pack(side="right")
        tk.Button(btns, text="Iniciar Processamento", bg="#2196F3", fg="white", command=self.iniciar).grid(row=0, column=0, padx=4)
        tk.Button(btns, text="Cancelar", bg="#f44336", fg="white", command=self.cancelar).grid(row=0, column=1, padx=4)
        tk.Button(btns, text="Exportar Relatório (Excel)", bg="#FF9800", command=self.abrir_relatorio).grid(row=0, column=2, padx=4)
        tk.Button(btns, text="Limpar Lista", command=self.limpar).grid(row=0, column=3, padx=4)

    def selecionar_origem(self):
        p = filedialog.askdirectory(title="Pasta de origem")
        if p:
            self.orig_folder.set(p)
            self.atualizar_lista()

    def selecionar_destino(self):
        p = filedialog.askdirectory(title="Pasta de destino (onde salvar PDFs)")
        if p:
            self.dest_folder.set(p)

    def atualizar_lista(self):
        self.tree.delete(*self.tree.get_children())
        self.registros.clear()
        pasta = self.orig_folder.get().strip()
        filtro = self.ext_filtro.get().strip().lower()
        if not pasta or not os.path.isdir(pasta):
            return
        arquivos = [f for f in os.listdir(pasta) if os.path.isfile(os.path.join(pasta, f))]
        arquivos.sort()
        for f in arquivos:
            if filtro and not f.lower().endswith(filtro):
                continue
            item = self.tree.insert("", "end", values=(f, "", "Aguardando"))
            rec = {
                "item": item,
                "antigo": f,
                "novo": "",
                "status": "Aguardando",
                "tipo": os.path.splitext(f)[1].lower(),
                "orig_path": os.path.join(pasta, f),
                "dest_path": "",
                "keyword": "",
                "timestamp": "",
                "mensagem": ""
            }
            self.registros.append(rec)
        self.progress['value'] = 0
        self.progress['maximum'] = max(1, len(self.registros))

    def iniciar(self):
        if not self.nome_base.get().strip():
            messagebox.showwarning("Aviso", "Informe o Nome base (obrigatório).")
            return
        if not self.orig_folder.get().strip() or not os.path.isdir(self.orig_folder.get().strip()):
            messagebox.showwarning("Aviso", "Selecione uma pasta de origem válida.")
            return
        if not self.dest_folder.get().strip() or not os.path.isdir(self.dest_folder.get().strip()):
            # pede pasta destino se não selecionada
            p = filedialog.askdirectory(title="Selecione a pasta de destino para salvar os PDFs")
            if not p:
                messagebox.showwarning("Aviso", "Pasta de destino necessária.")
                return
            self.dest_folder.set(p)

        if self.thread and self.thread.is_alive():
            messagebox.showinfo("Atenção", "Processamento já em execução.")
            return

        # prepara backup folder se necessário (inside origem)
        if self.backup_var.get():
            ensure_dir(os.path.join(self.orig_folder.get(), "BACKUP"))

        self.cancel_flag = False
        self.progress['value'] = 0
        self.progress['maximum'] = max(1, len(self.registros))
        self.thread = threading.Thread(target=self._processar_thread, daemon=True)
        self.thread.start()

    def cancelar(self):
        if self.thread and self.thread.is_alive():
            self.cancel_flag = True
            messagebox.showinfo("Cancelamento", "Pedido de cancelamento enviado. Aguardando conclusão da etapa atual.")
        else:
            messagebox.showinfo("Info", "Nenhum processamento em execução.")

    def limpar(self):
        if self.thread and self.thread.is_alive():
            messagebox.showwarning("Atenção", "Não é possível limpar durante o processamento.")
            return
        self.tree.delete(*self.tree.get_children())
        self.registros.clear()
        self.progress['value'] = 0

    def _processar_thread(self):
        origem = self.orig_folder.get().strip()
        destino = self.dest_folder.get().strip()
        nome_base = self.nome_base.get().strip()
        palavra_chave_param = self.palavra_chave.get().strip()
        backup = self.backup_var.get()

        total = len(self.registros)
        idx = 0
        for rec in self.registros:
            if self.cancel_flag:
                break
            idx += 1
            item = rec["item"]
            antigo = rec["antigo"]
            path_origem = rec["orig_path"]
            tipo = rec["tipo"]

            # update UI status
            self._update_tree(item, novo=rec.get("novo", ""), status="Processando")

            # extrair chave (se aplicável)
            try:
                if tipo == ".pdf":
                    chave = extrair_chave_pdf(path_origem, palavra_chave_param)
                elif tipo in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"):
                    chave = extrair_chave_imagem(path_origem, palavra_chave_param)
                elif tipo == ".txt":
                    chave = extrair_chave_txt(path_origem, palavra_chave_param)
                elif tipo == ".docx":
                    chave = extrair_chave_docx(path_origem, palavra_chave_param)
                else:
                    chave = None
            except Exception as e:
                chave = None

            # monta nome final: somente nome_base, se houver chave adiciona _chave
            if chave:
                novo_base = f"{nome_base}_{chave}"
            else:
                novo_base = nome_base

            novo_nome = novo_base + ".pdf"
            caminho_destino = os.path.join(destino, novo_nome)
            caminho_destino = generate_unique_path(caminho_destino)

            # backup original
            if backup:
                try:
                    shutil.copy2(path_origem, os.path.join(origem, "BACKUP", antigo))
                except Exception as e:
                    print("Backup falhou:", e)

            # converter
            sucesso = converter_para_pdf(path_origem, caminho_destino)
            estado = "Concluído" if sucesso else "Erro"
            msg = "" if sucesso else "Falha conversão"

            # grava dados no registro
            rec["novo"] = os.path.basename(caminho_destino) if sucesso else ""
            rec["status"] = estado
            rec["dest_path"] = caminho_destino if sucesso else ""
            rec["keyword"] = chave or ""
            rec["timestamp"] = datetime.now().isoformat(sep=' ', timespec='seconds')
            rec["mensagem"] = msg

            # atualiza UI
            self._update_tree(item, novo=rec["novo"], status=rec["status"])
            # update progress
            self.progress['value'] = idx
            time.sleep(0.05)

        # fim
        if self.cancel_flag:
            messagebox.showinfo("Processamento", "Processamento cancelado.")
        else:
            messagebox.showinfo("Processamento", "Processamento finalizado.")
        self.cancel_flag = False

    def _update_tree(self, item, novo=None, status=None):
        # thread-safe UI update
        def ui():
            vals = self.tree.item(item, "values")
            antigo = vals[0] if vals else ""
            novo_val = novo if novo is not None else (vals[1] if vals else "")
            status_val = status if status is not None else (vals[2] if vals else "")
            self.tree.item(item, values=(antigo, novo_val, status_val))
        self.root.after(0, ui)

    # ---------- relatório ----------
    def abrir_relatorio(self):
        win = tk.Toplevel(self.root)
        win.title("Opções do Relatório (Excel)")
        win.geometry("420x380")
        tk.Label(win, text="Selecione colunas a incluir:").pack(anchor="w", padx=10, pady=6)

        options = [
            ("Nome Original", "antigo"),
            ("Novo Nome", "novo"),
            ("Status", "status"),
            ("Tipo de Arquivo", "tipo"),
            ("Caminho Origem", "orig_path"),
            ("Caminho Destino", "dest_path"),
            ("Timestamp", "timestamp"),
            ("Palavra-chave", "keyword"),
            ("Mensagem", "mensagem")
        ]
        vars_map = {}
        for label, key in options:
            v = tk.BooleanVar(value=(key in ("antigo", "novo", "status")))
            chk = tk.Checkbutton(win, text=label, variable=v)
            chk.pack(anchor="w", padx=14)
            vars_map[key] = v

        def gerar():
            cols = [k for k, v in vars_map.items() if v.get()]
            if not cols:
                messagebox.showwarning("Atenção", "Escolha ao menos uma coluna.")
                return
            caminho = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                   filetypes=[("Excel", "*.xlsx")],
                                                   title="Salvar relatório como")
            if not caminho:
                return
            try:
                wb = openpyxl.Workbook()
                ws = wb.active
                # cabeçalho legível
                header = []
                for k in cols:
                    header.append(next(label for (label, key) in options if key == k))
                ws.append(header)
                for rec in self.registros:
                    row = [rec.get(k, "") for k in cols]
                    ws.append(row)
                wb.save(caminho)
                messagebox.showinfo("Relatório", f"Relatório salvo em: {caminho}")
                win.destroy()
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao gerar relatório: {e}")

        tk.Button(win, text="Gerar Relatório", bg="#4CAF50", fg="white", command=gerar).pack(pady=12)
        tk.Button(win, text="Fechar", command=win.destroy).pack()

# run
def main():
    root = tk.Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()

# ===== PyInstaller example =====
# pip install pyinstaller
# pyinstaller --onefile --noconsole --add-data "C:\\poppler-25.11.0\\Library\\bin;poppler_bin" --add-data "C:\\Program Files\\Tesseract-OCR\\tessdata;tessdata" app.py
#
# Ajuste paths conforme instalado no seu sistema.
